from typing import Dict, Any, Tuple
from docx import Document

import os
import shutil

from backend import replace
from config import TEMPLATES_PATH, OUTPUT_PATH

def parse_doc(input_path: str, output_path: str, data: Tuple[Dict, Dict]):
    fields, series = data
        
    doc = Document(input_path)
    
    for paragraph in doc.paragraphs:
        for placeholder, replacement in fields.items():
            replace.replace_placeholder_in_paragraph(paragraph, placeholder, replacement)
            
        for placeholder, replacement in series.items():
            replace.replace_series_in_paragraph(doc, paragraph, placeholder, replacement)
    
    for table in doc.tables:
        for placeholder, replacement in fields.items():
            replace.replace_placeholder_in_table(table, placeholder, replacement)
        
        for placeholder, replacement in series.items():
            replace.replace_series_in_table(table, placeholder, replacement)
    
    doc.save(output_path)

def create_unique_directory(base_path, dir_name):
    new_dir_path = os.path.join(base_path, dir_name)
    
    if not os.path.exists(new_dir_path):
        os.makedirs(new_dir_path)
        print(f"Directory created: {new_dir_path}")
        return new_dir_path + os.sep
    
    counter = 1
    while True:
        numbered_dir_name = f"{dir_name} ({counter})"
        new_dir_path = os.path.join(base_path, numbered_dir_name)
        if not os.path.exists(new_dir_path):
            os.makedirs(new_dir_path)
            print(f"Directory created: {new_dir_path}")
            return new_dir_path + os.sep
        counter += 1
    

def parse_documents(data: Tuple[Dict, Dict], path: str, export_window):
    dir_path = create_unique_directory(path, data[0]["[Nume complet]"])
    
    input_paths = [TEMPLATES_PATH + filename for filename in os.listdir(TEMPLATES_PATH)]
    output_paths = [dir_path + filename for filename in os.listdir(TEMPLATES_PATH)]
    
    for index, (input_path, output_path) in enumerate( zip(input_paths, output_paths), start=1 ):
        if input_path[-1] != '#':
            output_path = output_path.replace("[Nume complet]", data[0]["[Nume complet]"])
            
            parse_doc(input_path, output_path, data)
            
            export_window.update(index, len(input_paths), os.path.basename(output_path))
        
        