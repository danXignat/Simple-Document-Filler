from docx import Document, table
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List
import lxml

def replace_placeholder_in_paragraph(paragraph, placeholder: str, replacement: str):
    if placeholder in paragraph.text:
        paragraph.text = paragraph.text.replace(placeholder, str(replacement))

def replace_placeholder_in_table(table, placeholder: str, replacement: str):
    for row in table.rows:
        for cell in row.cells:
            if placeholder in cell.text:
                cell.text = cell.text.replace(placeholder, str(replacement))

def replace_series_in_table(table, placeholder: str, series: List):
    for row in table.rows:
        for cell in row.cells:
            if placeholder in cell.text:
                series_text = "\n".join(serie for serie in series)
                
                cell.text = cell.text.replace(placeholder, series_text)

def replace_series_in_paragraph(doc, paragraph, placeholder: str, series: List):
    if placeholder not in paragraph.text:
        return
    
    if len(series) == 1:
        paragraph.text = paragraph.text.replace(placeholder, series[0])
    else:
        paragraph.text = paragraph.text.replace(placeholder, "")
        
        table = doc.add_table(rows = len(series), cols = 1)
        
        nsmap=table._element[0].nsmap
        searchtag='{%s}tblPr' % nsmap['w']
        mytag='{%s}tblInd' % nsmap['w']
        myw='{%s}w' % nsmap['w']
        mytype='{%s}type' % nsmap['w']
        for elt in table._element:
            if elt.tag == searchtag:
                myelt=lxml.etree.Element(mytag)
                myelt.set(myw,'1000')
                myelt.set(mytype,'dxa')
                myelt=elt.append(myelt)
        
        table.autofit = False 
        table.allow_autofit = False
        table.columns[0].width = Inches(1.0)
        table.rows[0].cells[0].width = Inches(1.0)
        
        for i, text in enumerate(series):
            cell = table.cell(i, 0)
            cell.text = text
        
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            
        paragraph._element.addnext(table._tbl)