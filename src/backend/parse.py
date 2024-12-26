from typing import Dict, Any, Tuple
from docx import Document

import os
import shutil

from backend import replace
from config import TEMPLATES_PATH, OUTPUT_PATH

def parse_doc(input_path: str, output_path: str, data: Tuple[Dict, Dict]):
    fields, series = data
        
    doc = Document(input_path)
    
    for paragraph in doc.paragraphs:
        for placeholder, replacement in fields.items():
            replace.replace_placeholder_in_paragraph(paragraph, placeholder, replacement)
            
        for placeholder, replacement in series.items():
            replace.replace_series_in_paragraph(doc, paragraph, placeholder, replacement)
    
    for table in doc.tables:
        for placeholder, replacement in fields.items():
            replace.replace_placeholder_in_table(table, placeholder, replacement)
        
        for placeholder, replacement in series.items():
            replace.replace_series_in_table(table, placeholder, replacement)
    
    doc.save(output_path)
    
def parse_documents(data: Tuple[Dict, Dict], path: str, export_window):
    input_paths = [TEMPLATES_PATH + filename for filename in os.listdir(TEMPLATES_PATH)]
    output_paths = [path + filename for filename in os.listdir(TEMPLATES_PATH)]
    
    # shutil.rmtree(OUTPUT_PATH)
    # os.mkdir(OUTPUT_PATH)
    
    for index, (input_path, output_path) in enumerate( zip(input_paths, output_paths), start=1 ):
        if input_path[-1] != '#':
            output_path = output_path.replace("[Nume complet]", data[0]["[Nume complet]"])
            
            parse_doc(input_path, output_path, data)
            
            export_window.update(index, len(input_paths), os.path.basename(output_path))
        
        