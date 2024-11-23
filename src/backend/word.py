from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any

import os

from config import TEMPLATES_PATH, OUTPUT_PATH

def create_table(doc : Document, panels : List[str]):
    table = doc.add_table(rows = 1, cols = 1)
    
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.5)
    table.rows[0].cells[0].width = Inches(1.5)
    
    header_cell = table.rows[0].cells[0]
    header_cell.text = "Serie Panouri"
    
    for panel in panels:
        row = table.add_row()
        cell = row.cells[0]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(panel)
        run.bold = True
            
    return table
    
def replace_placeholders(data : Dict[str, Any], doc_name : str, save_path : str):
    doc = Document(TEMPLATES_PATH + doc_name)
    
    for paragraph in doc.paragraphs:
        
        for key in data.keys():
            placeholder =  f"[{key}]"
            
            if placeholder in paragraph.text and placeholder == "[panouri]":
                paragraph.text = paragraph.text.replace(placeholder, "")
                
                table = create_table(doc, data["panouri"])
                
                paragraph._element.addnext(table._element)
                
                continue
            
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, data[key])
                
    new_name = doc_name.replace("[Nume]", data["Nume"])
    doc.save(save_path + new_name)
    

def submit(data : Dict[str, Any], save_path : str):
    filenames = os.listdir(TEMPLATES_PATH)
    
    for filename in filenames:
        replace_placeholders(data, filename, save_path)    
    
