from config import TEMPLATES_PATH
from docx import Document, table
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import os
from icecream import ic

column_headers = [
        "Denumire",
        "Panou solar",
        "Invertor",
        "Smart Meter",
        "Panou distributie",
        "Sistem prindere",
        "Accesorii"
    ]
    
row_labels = [
        "UM",
        "Cantitate",
        "Putere minima a panoului",
        "Tehnologie panouri",
        "Rama panou",
        "Conectare",
        "Eficienta",
        "Grad protectie",
        "Interval de temper de functionare",
        "Garantie",
        "Garantie eficienta",
        "Putere nominala instalata",
        "MPPT",
        "Iesire",
        "Frecventa",
        "Umiditate max 95%",
        "Serii"
    ]

def replace_placeholder_with_table(doc_path, output_path, table_data):
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
        if "[big_table]" in paragraph.text:
            placeholder_paragraph = paragraph
            break

    rows, cols = len(table_data), len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)

    for j, text in enumerate(column_headers):
        cell = table.cell(0, j)
        cell.text = text
    
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    for i, text in enumerate(row_labels, start=1):
        cell = table.cell(i, 0)
        cell.text = text
    
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    multiple_series = [i for i in range(16)]
    
    cell = table.cell(17, 1)
    cell.text = ""
    
    for series in multiple_series:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(str(series))
        run.font.size = Pt(10)
    
    # for i, row_data in enumerate(table_data):
    #     for j, cell_text in enumerate(row_data):
    #         cell = table.cell(i, j)
    #         cell.text = str(cell_text)

    #         cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #         run = cell.paragraphs[0].runs[0]
    #         run.font.size = Pt(10)

    placeholder_paragraph.clear()  # Clear the placeholder text
    placeholder_paragraph._element.addnext(table._tbl)

    try:
        os.remove(output_path)  # Deletes the file
        print(f"File '{output_path}' has been deleted successfully.")
    except FileNotFoundError:
        print(f"Error: The file '{output_path}' does not exist.")
    except PermissionError:
        print(f"Error: You do not have permission to delete '{output_path}'.")
    except Exception as e:
        print(f"An error occurred: {e}")
        
    doc.save(output_path)


# Example Usage
if __name__ == "__main__":
    doc_path = TEMPLATES_PATH + "Proces verbal de predare primire [Nume].docx"
    output_path = "Proces verbal de predare primire Output.docx"

    table = []
    table.append(column_headers)
    
    for row_label in row_labels:
        row = [row_label]
        
        row += ['-' for _ in range(len(column_headers) - 1)]
        
        table.append(row)
    
    ic(table)

    replace_placeholder_with_table(doc_path, output_path, table)

